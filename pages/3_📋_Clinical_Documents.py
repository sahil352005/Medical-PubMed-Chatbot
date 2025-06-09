import streamlit as st
from mistral_chains import answer_user_query
import pandas as pd
from datetime import datetime
import json
import os
from pathlib import Path
import base64
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

st.set_page_config(page_title="Clinical Documents", layout="wide")
st.title("📋 Clinical Documents")

# Check if articles are available in session state
if "articles" not in st.session_state or st.session_state.articles is None:
    st.warning("⚠ Please process articles in the Evidence Analysis page first.")
    st.stop()

articles = st.session_state.articles

# Initialize document data in session state if not exists
for key in ['protocol_intro', 'csr_discussion']:
    if key not in st.session_state:
        st.session_state[key] = None

# Create tabs for different documents
doc_tab1, doc_tab2 = st.tabs([
    "📄 Clinical Trial Protocol",
    "📑 Clinical Study Report"
])

# Generate documents if not already in session state
if st.session_state.protocol_intro is None:
    with st.spinner("Generating Clinical Trial Protocol Introduction..."):
        st.session_state.protocol_intro = answer_user_query("""
        Generate the Introduction section for a Clinical Trial Protocol in ICH M11 format with these subheadings:
        1. Background and Rationale
        2. Current Treatment Landscape
        3. Unmet Need
        4. Justification for Investigating Study Drug

        Use the provided articles as evidence.
        """, articles, "")

if st.session_state.csr_discussion is None:
    with st.spinner("Generating Clinical Study Report Discussion..."):
        st.session_state.csr_discussion = answer_user_query("""
        Generate the Discussion section of a Clinical Study Report in TransCelerate format covering:
        1. Interpretation of Efficacy Outcomes
        2. Interpretation of Safety Outcomes
        3. Consistency of Findings
        4. Study Limitations
        5. Implications for Future Clinical Trials

        Use the provided articles as evidence.
        """, articles, "")

# Display documents in tabs
with doc_tab1:
    st.markdown("### Clinical Trial Protocol Introduction (ICH M11 Format)")
    st.markdown(st.session_state.protocol_intro)
    
    # Download button for protocol
    protocol_text = f"# Clinical Trial Protocol Introduction\n\n{st.session_state.protocol_intro}"
    st.download_button(
        label="📥 Download Protocol Introduction",
        data=protocol_text,
        file_name="protocol_introduction.md",
        mime="text/markdown"
    )

with doc_tab2:
    st.markdown("### Clinical Study Report Discussion (TransCelerate Format)")
    st.markdown(st.session_state.csr_discussion)
    
    # Download button for CSR
    csr_text = f"# Clinical Study Report Discussion\n\n{st.session_state.csr_discussion}"
    st.download_button(
        label="📥 Download CSR Discussion",
        data=csr_text,
        file_name="csr_discussion.md",
        mime="text/markdown"
    )

# Add tips for better document generation
with st.expander("💡 Tips for Better Documents"):
    st.markdown("""
    To get better document generation results:
    1. Ensure articles are relevant to your research question
    2. Include a mix of recent and landmark studies
    3. Make sure articles contain detailed methodology and results
    4. Consider including both positive and negative outcomes
    """)

def create_clinical_protocol(articles_data, patient_info):
    """Create a clinical protocol document based on evidence and patient info."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Clinical Protocol', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Patient Information
    doc.add_heading('Patient Information', level=1)
    for key, value in patient_info.items():
        doc.add_paragraph(f"{key}: {value}")
    
    # Evidence Summary
    doc.add_heading('Evidence Summary', level=1)
    doc.add_paragraph(articles_data.get('evidence_summary', 'No evidence summary available.'))
    
    # Clinical Recommendations
    doc.add_heading('Clinical Recommendations', level=1)
    doc.add_paragraph(articles_data.get('clinical_recommendations', 'No clinical recommendations available.'))
    
    # Implementation Plan
    doc.add_heading('Implementation Plan', level=1)
    doc.add_paragraph(articles_data.get('implementation_plan', 'No implementation plan available.'))
    
    # Monitoring and Follow-up
    doc.add_heading('Monitoring and Follow-up', level=1)
    doc.add_paragraph(articles_data.get('monitoring_plan', 'No monitoring plan available.'))
    
    # References
    doc.add_heading('References', level=1)
    for article in articles_data.get('articles', []):
        p = doc.add_paragraph()
        p.add_run(f"{article.get('title', 'No title')} ").bold = True
        p.add_run(f"({article.get('authors', 'No authors')}, {article.get('year', 'No year')})")
        p.add_run(f"\n{article.get('abstract', 'No abstract available.')}")
    
    return doc

def get_download_link(doc, filename):
    """Generate a download link for the document."""
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    b64 = base64.b64encode(bio.read()).decode()
    href = f'data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}'
    return href

def get_protocol_preview(doc):
    """Generate a text preview of the protocol."""
    preview = []
    for para in doc.paragraphs:
        if para.text.strip():
            if para.style.name.startswith('Heading'):
                preview.append(f"\n{'#' * int(para.style.name[-1])} {para.text}\n")
            else:
                preview.append(para.text)
    return "\n".join(preview)

def main():
    st.title("📋 Clinical Documents")
    
    # Check if we have the required data
    if 'articles_data' not in st.session_state:
        st.warning("Please analyze articles in the Evidence Analysis page first.")
        return
    
    articles_data = st.session_state.articles_data
    
    # Patient Information Form
    st.header("Patient Information")
    col1, col2 = st.columns(2)
    
    with col1:
        patient_id = st.text_input("Patient ID")
        age = st.number_input("Age", min_value=0, max_value=120)
        gender = st.selectbox("Gender", ["Male", "Female", "Other"])
    
    with col2:
        diagnosis = st.text_area("Primary Diagnosis")
        comorbidities = st.text_area("Comorbidities")
        current_medications = st.text_area("Current Medications")
    
    patient_info = {
        "Patient ID": patient_id,
        "Age": age,
        "Gender": gender,
        "Primary Diagnosis": diagnosis,
        "Comorbidities": comorbidities,
        "Current Medications": current_medications
    }
    
    # Generate Protocol Button
    if st.button("Generate Clinical Protocol"):
        with st.spinner("Generating clinical protocol..."):
            # Create the document
            doc = create_clinical_protocol(articles_data, patient_info)
            
            # Store the document in session state
            st.session_state.clinical_protocol = doc
            
            # Show success message
            st.success("Clinical protocol generated successfully!")
            
            # Show preview in an expander
            with st.expander("Preview Protocol", expanded=True):
                preview_text = get_protocol_preview(doc)
                st.text_area("Protocol Preview", preview_text, height=400)
            
            # Download button
            href = get_download_link(doc, "clinical_protocol.docx")
            st.markdown(f'<a href="{href}" download="clinical_protocol.docx">Download Clinical Protocol</a>', unsafe_allow_html=True)

if __name__== "_main_":
    main()