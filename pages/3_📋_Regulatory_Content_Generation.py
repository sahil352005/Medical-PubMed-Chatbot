# HARDCODED_CSR_DISCUSSION = """
# Clinical Study Report – Discussion (TransCelerate Format)

# 1. Interpretation of Efficacy Outcomes
# The studies consistently indicate that vonoprazan demonstrates superior or comparable efficacy to traditional proton pump inhibitors (PPIs) in treating peptic and duodenal ulcers. Meng et al. (2023) further supports that PPIs are more effective than H2 receptor antagonists in initial treatment settings. Ai et al. (2024) show that co-administration of PPIs with clopidogrel in post-PCI patients is generally effective without significant compromise in cardiovascular outcomes.

# 2. Interpretation of Safety Outcomes
# Across all three articles, the safety profiles of PPIs and vonoprazan are generally well-tolerated, with low rates of serious adverse events. Vonoprazan showed a favorable safety profile, potentially offering a safer alternative to long-term PPI use. The combination therapy with clopidogrel also did not significantly increase bleeding risk, although individual patient factors should be considered.

# 3. Consistency of Findings
# There is consistency in the findings across multiple randomized controlled trials that support the efficacy and safety of acid suppression therapies. Vonoprazan consistently ranks high in network meta-analyses for both efficacy and safety, with PPIs also performing reliably. This consistency enhances confidence in integrating these treatments into routine care.

# 4. Study Limitations
# While the network meta-analyses provide strong comparative insights, heterogeneity in study design, treatment duration, and patient populations may limit direct generalizability. Furthermore, long-term safety data for vonoprazan remains limited.

# 5. Implications for Future Clinical Trials
# Future trials should focus on head-to-head comparisons in diverse populations, with extended follow-up to evaluate long-term outcomes, particularly for vonoprazan. Evaluating outcomes in high-risk subgroups, such as patients on dual antiplatelet therapy, is also warranted.
# """

# import streamlit as st
# from mistral_chains import answer_user_query
# from docx import Document
# from docx.shared import Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# import base64
# import io

# st.set_page_config(page_title="Regulatory Content Generation", layout="wide")

# st.image("https://s3ktech.ai/wp-content/uploads/2025/03/S3Ktech-Logo.png", width=140)

# st.title("📋 Regulatory Content Generation")

# if "articles" not in st.session_state or st.session_state.articles is None:
#     st.warning("⚠ Please process articles in the Evidence Analysis page first.")
#     st.stop()

# articles = st.session_state.articles

# # Tabs for Protocol and CSR
# doc_tab1, doc_tab2 = st.tabs(["📄 Clinical Trial Protocol", "📑 Clinical Study Report"])

# # Generate Introduction
# if "protocol_intro" not in st.session_state:
#     with st.spinner("Generating Clinical Trial Protocol Introduction..."):
#         st.session_state.protocol_intro = answer_user_query("""
#         Generate the Introduction section for a Clinical Trial Protocol in ICH M11 format with these subheadings:
#         1. Background and Rationale
#         2. Current Treatment Landscape
#         3. Unmet Need
#         4. Justification for Investigating Study Drug
#         """, articles, "")

# # Generate CSR Discussion
# if "csr_discussion" not in st.session_state:
#     with st.spinner("Generating Clinical Study Report Discussion..."):
#         st.session_state.csr_discussion = answer_user_query("""
#         Generate the Discussion section of a Clinical Study Report in TransCelerate format covering:
#         1. Interpretation of Efficacy Outcomes
#         2. Interpretation of Safety Outcomes
#         3. Consistency of Findings
#         4. Study Limitations
#         5. Implications for Future Clinical Trials
#         """, articles, "")

# # 📄 Clinical Trial Protocol Tab
# def get_protocol_preview(doc):
#     preview = []
#     for para in doc.paragraphs:
#         if para.text.strip():
#             preview.append(para.text)
#     return "\n".join(preview)

# with doc_tab1:
#     st.markdown("""
#         <style>
#         .protocol-preview-card {
#             background: #f8f9fa;
#             border-radius: 10px;
#             border: 1px solid #e9ecef;
#             padding: 24px 18px 18px 18px;
#             margin-bottom: 24px;
#             font-family: 'Segoe UI', 'Arial', sans-serif;
#             font-size: 1.08rem;
#             color: #222;
#             box-shadow: 0 2px 8px rgba(0,0,0,0.04);
#             line-height: 1.7;
#             max-height: 400px;
#             overflow-y: auto;
#         }
#         .protocol-preview-title {
#             font-size: 1.15rem;
#             font-weight: 600;
#             color: #b22222;
#             margin-bottom: 10px;
#             letter-spacing: 0.5px;
#         }
#         </style>
#     """, unsafe_allow_html=True)
    
#     st.markdown("<span style='font-size:1.2rem;font-weight:600;'>📄 Full Clinical Protocol Document</span>", unsafe_allow_html=True)

#     if st.button("Generate Full Clinical Protocol Document"):
#         with st.spinner("Generating full protocol..."):

#             # Generate 100-word per-article summaries
#             per_article_summaries = []
#             for i, art in enumerate(articles, 1):
#                 summary = answer_user_query(
#                     "Generate a 100-word structured summary for this article including Background, Methodology, Key Findings, and Conclusion",
#                     [art],
#                     art['abstract'],
#                     forced_tasks=["per_article_summary"]
#                 )
#                 per_article_summaries.append(f"📝 Article {i}:\n{summary.strip()}")

#             # Create DOCX
#             doc = Document()
#             title = doc.add_heading('Clinical Protocol', 0)
#             title.alignment = WD_ALIGN_PARAGRAPH.CENTER

#             doc.add_heading('Evidence Summary', level=1)
#             for summary in per_article_summaries:
#                 doc.add_paragraph(summary)

#             doc.add_heading('Clinical Recommendations', level=1)
#             doc.add_paragraph("Recommendations based on evidence will be provided here.")

#             doc.add_heading('Implementation Plan', level=1)
#             doc.add_paragraph("Implementation steps for applying this protocol in a clinical setting.")

#             doc.add_heading('Monitoring and Follow-up', level=1)
#             doc.add_paragraph("Details of follow-up assessments, lab tests, and monitoring strategy.")

#             doc.add_heading('References', level=1)
#             for art in articles:
#                 ref = doc.add_paragraph()
#                 ref.add_run(f"{art.get('title', 'No title')} ").bold = True
#                 ref.add_run(f"({art.get('authors', 'No authors')}, {art.get('year', 'No year')})")

#             # Save to buffer
#             buffer = io.BytesIO()
#             doc.save(buffer)
#             buffer.seek(0)
#             b64 = base64.b64encode(buffer.read()).decode()
#             href = f'data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}'

#             # Preview
#             preview_text = get_protocol_preview(doc)
#             st.markdown("<div class='protocol-preview-title'>Preview</div>", unsafe_allow_html=True)
#             st.markdown(f"<div class='protocol-preview-card'>{preview_text.replace(chr(10), '<br>')}</div>", unsafe_allow_html=True)

#             # Download link
#             st.markdown(f'<a href="{href}" download="clinical_protocol.docx"><b>📥 Download Clinical Protocol</b></a>', unsafe_allow_html=True)

# # 📑 Clinical Study Report Tab
# with doc_tab2:
#     st.markdown("""
#         <style>
#         .csr-preview-card {
#             background: #f8f9fa;
#             border-radius: 10px;
#             border: 1px solid #e9ecef;
#             padding: 24px 18px 18px 18px;
#             margin-bottom: 24px;
#             font-family: 'Segoe UI', 'Arial', sans-serif;
#             font-size: 1.08rem;
#             color: #222;
#             box-shadow: 0 2px 8px rgba(0,0,0,0.04);
#             line-height: 1.7;
#             max-height: 400px;
#             overflow-y: auto;
#         }
#         .csr-preview-title {
#             font-size: 1.15rem;
#             font-weight: 600;
#             color: #b22222;
#             margin-bottom: 10px;
#             letter-spacing: 0.5px;
#         }
#         </style>
#     """, unsafe_allow_html=True)

#     st.markdown("<span style='font-size:1.2rem;font-weight:600;'>CSR Discussion (TransCelerate Format)</span>", unsafe_allow_html=True)

#     # Use your hardcoded CSR discussion text or st.session_state.csr_discussion if you want dynamic
#     csr_text = HARDCODED_CSR_DISCUSSION

#     # Preview
#     st.markdown("<div class='csr-preview-title'>Preview</div>", unsafe_allow_html=True)
#     st.markdown(f"<div class='csr-preview-card'>{csr_text.replace(chr(10), '<br>')}</div>", unsafe_allow_html=True)

#     # Download as Markdown
#     st.download_button(
#         label="📥 Download CSR Discussion (Markdown)",
#         data=f"# Clinical Study Report Discussion\n\n{csr_text}",
#         file_name="csr_discussion.md",
#         mime="text/markdown"
#     )

#     # Download as DOCX
#     def generate_csr_docx(text):
#         doc = Document()
#         title = doc.add_heading('Clinical Study Report - Discussion', 0)
#         title.alignment = WD_ALIGN_PARAGRAPH.CENTER

#         for line in text.strip().split('\n\n'):
#             if line.startswith('###'):
#                 doc.add_heading(line.replace('###', '').strip(), level=3)
#             elif line.startswith('####'):
#                 doc.add_heading(line.replace('####', '').strip(), level=4)
#             else:
#                 doc.add_paragraph(line.strip())
#         return doc

#     if st.button("Download CSR Discussion as DOCX"):
#         doc = generate_csr_docx(csr_text)
#         buffer = io.BytesIO()
#         doc.save(buffer)
#         buffer.seek(0)
#         b64 = base64.b64encode(buffer.read()).decode()
#         href = f'data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}'
#         st.markdown(f'<a href="{href}" download="csr_discussion.docx"><b>📥 Download CSR Discussion</b></a>', unsafe_allow_html=True)

import streamlit as st
from mistral_chains import answer_user_query
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
import io

st.set_page_config(page_title="Clinical Protocol Generator", layout="wide")

# --- Custom CSS ---
st.markdown("""
    <style>
    body {
        font-family: 'Segoe UI', sans-serif;
        background-color: #f7f9fc;
    }
    .main-title {
        text-align: center;
        font-size: 2.2rem;
        font-weight: 700;
        color: black;
        margin-bottom: 10px;
    }
    .section-title {
        font-size: 1.5rem;
        color: #343a40;
        font-weight: 600;
        margin: 20px 0 10px;
    }
    .preview-box {
        background-color: #ffffff;
        border-radius: 12px;
        padding: 20px;
        border: 1px solid #dee2e6;
        box-shadow: 0px 4px 12px rgba(0,0,0,0.06);
        max-height: 400px;
        overflow-y: auto;
        margin-bottom: 20px;
        line-height: 1.7;
        color: #212529;
    }
    .stButton>button {
        background-color: #2563eb;
        color: white;
        border-radius: 8px;
        padding: 0.6em 1.2em;
        font-size: 1rem;
        font-weight: 600;
        border: none;
        box-shadow: 0px 3px 6px rgba(0,0,0,0.1);
    }
    .stButton>button:hover {
        background-color: #1d4ed8;
        transition: 0.3s ease-in-out;
    }
    .download-link {
        font-weight: 600;
        color: #2563eb;
        font-size: 1.05rem;
    }
    .logo {
        display: block;
        margin-left: auto;
        margin-right: auto;
        width: 140px;
    }
    </style>
""", unsafe_allow_html=True)

# --- Page Elements ---
st.image("https://s3ktech.ai/wp-content/uploads/2025/03/S3Ktech-Logo.png", width=140)
st.markdown("<div class='main-title'>📋 Clinical Protocol Generator</div>", unsafe_allow_html=True)

if "articles" not in st.session_state or st.session_state.articles is None:
    st.warning("⚠ Please process articles in the Evidence Analysis page first.")
    st.stop()

articles = st.session_state.articles

# --- Generate Introduction if not already done ---
if "protocol_intro" not in st.session_state:
    with st.spinner("🔄 Generating Protocol Introduction..."):
        st.session_state.protocol_intro = answer_user_query("""
        Generate the Introduction section for a Clinical Trial Protocol in ICH M11 format with these subheadings:
        1. Background and Rationale
        2. Current Treatment Landscape
        3. Unmet Need
        4. Justification for Investigating Study Drug
        """, articles, "")

# --- Generate Clinical Protocol ---
def get_protocol_preview(doc):
    preview = []
    for para in doc.paragraphs:
        if para.text.strip():
            preview.append(para.text)
    return "\n".join(preview)

st.markdown("<div class='section-title'>📄 Full Clinical Protocol Document</div>", unsafe_allow_html=True)

if st.button("Generate Full Protocol Document"):
    with st.spinner("🛠 Generating full protocol document..."):

        per_article_summaries = []
        for i, art in enumerate(articles, 1):
            summary = answer_user_query(
                "Generate a 100-word structured summary for this article including Background, Methodology, Key Findings, and Conclusion",
                [art],
                art['abstract'],
                forced_tasks=["per_article_summary"]
            )
            per_article_summaries.append(f"📝 Article {i}:\n{summary.strip()}")

        # Create the DOCX document
        doc = Document()
        title = doc.add_heading('Clinical Protocol', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('Evidence Summary', level=1)
        for summary in per_article_summaries:
            doc.add_paragraph(summary)

        doc.add_heading('Clinical Recommendations', level=1)
        doc.add_paragraph("Recommendations based on evidence will be provided here.")

        doc.add_heading('Implementation Plan', level=1)
        doc.add_paragraph("Implementation steps for applying this protocol in a clinical setting.")

        doc.add_heading('Monitoring and Follow-up', level=1)
        doc.add_paragraph("Details of follow-up assessments, lab tests, and monitoring strategy.")

        doc.add_heading('References', level=1)
        for art in articles:
            ref = doc.add_paragraph()
            ref.add_run(f"{art.get('title', 'No title')} ").bold = True
            ref.add_run(f"({art.get('authors', 'No authors')}, {art.get('year', 'No year')})")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        b64 = base64.b64encode(buffer.read()).decode()
        href = f'data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}'

        # Preview Section
        preview_text = get_protocol_preview(doc)
        st.markdown("<div class='section-title'>👀 Preview</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='preview-box'>{preview_text.replace(chr(10), '<br>')}</div>", unsafe_allow_html=True)

        # Download Link
        st.markdown(f'<a href="{href}" download="clinical_protocol.docx" class="download-link">📥 Download Clinical Protocol</a>', unsafe_allow_html=True)
